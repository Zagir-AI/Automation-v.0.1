"""
docs/gen_load_tables.py — ведомость нагрузок ГОСТ 21.613.

Три режима вывода:
  by_panel   — таблица по щитам (один щит = один блок)
  by_section — сводка по разделам смежников
  summary    — итоговая сводная таблица по объекту

Данные читаются из project["_results"] — не из vru напрямую.
"""

from __future__ import annotations
from pathlib import Path
from datetime import date

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False


# ── Вспомогательные функции ───────────────────────────────────────────────────

def _iter_panels(vru_results: dict):
    """Итератор по всем щитам из _results.vru.feeders."""
    for feeder in vru_results.get("feeders", []):
        for panel in feeder.get("panels", []):
            yield panel


def _iter_consumers(panel: dict):
    """Итератор по потребителям щита (пропускает резервные)."""
    for c in panel.get("consumers", []):
        if not c.get("reserve", False):
            yield c


def _panel_section(panel: dict) -> str:
    """Раздел щита по типу."""
    _type_to_section = {
        "lighting":       "ЭОМ",
        "power":          "ЭОМ",
        "heating":        "ОВ",
        "ventilation":    "ВК",
        "hvac":           "КВ",
        "technology":     "ТХ",
        "smoke_exhaust":  "ДУ",
        "firefighting":   "ПС",
        "outdoor_lighting": "ЭН",
    }
    return _type_to_section.get(panel.get("type", ""), "ТХ")


def _consumer_section(consumer: dict) -> str:
    return consumer.get("section", "ЭОМ")


# ── Текстовый fallback ────────────────────────────────────────────────────────

def _txt_load_table_by_panel(project: dict, docs_dir: Path) -> Path:
    vru = project.get("_results", {}).get("vru", {})
    proj_name = project.get("project", {}).get("name", "Объект")
    proj_code = project.get("project", {}).get("code", "")

    lines = []
    lines.append(f"ВЕДОМОСТЬ НАГРУЗОК — {proj_code} {proj_name}")
    lines.append(f"Дата: {date.today()}")
    lines.append("")

    col = "{:<6} {:<35} {:<10} {:<6} {:<8} {:<8} {:<8} {:<6} {:<18} {:<8}"
    hdr = col.format("Поз.", "Наименование", "Тип", "Кат.", "Pуст,кВт",
                     "Кс", "Pрасч,кВт", "cosφ", "Кабель", "АВ,А")
    sep = "-" * 120

    for panel in _iter_panels(vru):
        lines.append(f"\nЩИТ {panel['id']} — {panel['name']}")
        lines.append(sep)
        lines.append(hdr)
        lines.append(sep)

        for c in _iter_consumers(panel):
            cable = c.get("cable", {})
            cable_str = ""
            if cable:
                cable_str = (f"{cable.get('mark','')} "
                             f"{cable.get('cores','')}×"
                             f"{cable.get('section_mm2','')}мм²")
            br = c.get("breaker", {})
            br_str = str(br.get("rating", "—")) if br else "—"
            lines.append(col.format(
                c.get("id", ""),
                c.get("name", "")[:35],
                c.get("type", ""),
                str(c.get("category_pue", 3)),
                f"{c.get('power_kw', 0):.2f}",
                f"{c.get('demand_factor', 1):.2f}",
                f"{c.get('p_calc_kw', 0):.2f}",
                f"{c.get('cos_phi', 1):.2f}",
                cable_str[:18],
                br_str,
            ))

        lines.append(sep)
        lines.append(col.format(
            "", "ИТОГО по щиту:", "", "",
            f"{panel.get('p_installed_kw', 0):.2f}",
            "",
            f"{panel.get('p_calc_kw', 0):.2f}",
            f"{panel.get('cos_phi', 1):.2f}",
            "", ""
        ))

    # Итого по ВРУ
    lines.append("")
    lines.append("=" * 120)
    lines.append(f"ИТОГО по ВРУ:  "
                 f"Pуст={vru.get('p_installed_kw', 0):.2f} кВт  "
                 f"Pрасч={vru.get('p_calc_kw', 0):.2f} кВт  "
                 f"Iвру={vru.get('i_calc_a', 0):.2f} А  "
                 f"cosφ={vru.get('cos_phi', 1):.3f}")

    out = docs_dir / "load_table_by_panel.txt"
    out.write_text("\n".join(lines), encoding="utf-8")
    return out


def _txt_load_table_by_section(project: dict, docs_dir: Path) -> Path:
    vru = project.get("_results", {}).get("vru", {})
    proj_name = project.get("project", {}).get("name", "Объект")
    proj_code = project.get("project", {}).get("code", "")

    # Группируем по разделам
    sections: dict[str, dict] = {}
    for panel in _iter_panels(vru):
        sect = _panel_section(panel)
        if sect not in sections:
            sections[sect] = {"p_installed": 0.0, "p_calc": 0.0, "consumers": []}
        sections[sect]["p_installed"] += panel.get("p_installed_kw", 0)
        sections[sect]["p_calc"]      += panel.get("p_calc_kw", 0)
        for c in _iter_consumers(panel):
            c_sect = _consumer_section(c) or sect
            sections.setdefault(c_sect, {"p_installed": 0.0, "p_calc": 0.0, "consumers": []})
            sections[c_sect]["consumers"].append(c)

    lines = []
    lines.append(f"ВЕДОМОСТЬ НАГРУЗОК ПО РАЗДЕЛАМ — {proj_code} {proj_name}")
    lines.append(f"Дата: {date.today()}")
    lines.append("")

    for sect, data in sorted(sections.items()):
        p_inst = sum(c.get("power_kw", 0) for c in data["consumers"])
        p_calc = sum(c.get("p_calc_kw", 0) for c in data["consumers"])
        lines.append(f"Раздел {sect}:  Pуст={p_inst:.2f} кВт  Pрасч={p_calc:.2f} кВт  "
                     f"({len(data['consumers'])} потребителей)")

    total_p = vru.get("p_installed_kw", 0)
    total_q = vru.get("p_calc_kw", 0)
    lines.append("")
    lines.append(f"ИТОГО:  Pуст={total_p:.2f} кВт  Pрасч={total_q:.2f} кВт")

    out = docs_dir / "load_table_by_section.txt"
    out.write_text("\n".join(lines), encoding="utf-8")
    return out


def _txt_load_summary(project: dict, docs_dir: Path) -> Path:
    vru = project.get("_results", {}).get("vru", {})
    proj_name = project.get("project", {}).get("name", "Объект")
    proj_code = project.get("project", {}).get("code", "")
    comp = project.get("compensation", {})

    lines = []
    lines.append(f"СВОДНАЯ ТАБЛИЦА НАГРУЗОК — {proj_code} {proj_name}")
    lines.append(f"Дата: {date.today()}")
    lines.append("")
    lines.append(f"  Pуст       = {vru.get('p_installed_kw', 0):.2f} кВт")
    lines.append(f"  Pрасч      = {vru.get('p_calc_kw', 0):.2f} кВт")
    lines.append(f"  Sрасч      = {vru.get('s_calc_kva', 0):.2f} кВА")
    lines.append(f"  cosφ       = {vru.get('cos_phi', 1):.3f}")
    lines.append(f"  Iвру       = {vru.get('i_calc_a', 0):.2f} А")
    br = vru.get("breaker", {})
    if br:
        lines.append(f"  АВ ВРУ     = {br.get('type', '')}")
    cable = vru.get("cable", {})
    if cable:
        lines.append(f"  Кабель ВРУ = {cable.get('mark','')} "
                     f"{cable.get('cores','')}×{cable.get('section_mm2','')}мм²")

    if comp:
        lines.append("")
        lines.append(f"  КРМ необходима: {'Да' if comp.get('required') else 'Нет'}")
        if comp.get("selected_krm"):
            krm = comp["selected_krm"]
            lines.append(f"  Батарея КРМ: {krm['model']} ({krm['power_kvar']} кВАр)")

    lines.append("")
    lines.append(f"  Щитов: {sum(1 for _ in _iter_panels(vru))}")
    total_c = sum(1 for p in _iter_panels(vru) for _ in _iter_consumers(p))
    lines.append(f"  Потребителей: {total_c}")

    out = docs_dir / "load_summary.txt"
    out.write_text("\n".join(lines), encoding="utf-8")
    return out


# ── docx реализация ───────────────────────────────────────────────────────────

def _set_col_widths(table, widths_cm: list):
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if j < len(widths_cm):
                cell.width = Cm(widths_cm[j])


def _add_header_row(table, headers: list, font_size: int = 9):
    row = table.rows[0]
    for i, (cell, text) in enumerate(zip(row.cells, headers)):
        cell.text = text
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0] if para.runs else para.add_run(text)
        run.bold = True
        run.font.size = Pt(font_size)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _add_data_row(table, values: list, bold: bool = False, font_size: int = 8):
    row = table.add_row()
    for cell, val in zip(row.cells, values):
        cell.text = str(val)
        para = cell.paragraphs[0]
        run = para.runs[0] if para.runs else para.add_run(str(val))
        run.bold = bold
        run.font.size = Pt(font_size)
    return row


def _docx_load_table_by_panel(project: dict, docs_dir: Path) -> Path:
    vru = project.get("_results", {}).get("vru", {})
    proj_info = project.get("project", {})
    proj_name = proj_info.get("name", "Объект")
    proj_code = proj_info.get("code", "")

    doc = Document()
    doc.core_properties.author = "elec-system"

    # Стили
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)

    # Заголовок
    title = doc.add_heading("Ведомость нагрузок", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"{proj_code} — {proj_name}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Дата: {date.today()}  |  ГОСТ 21.613").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    headers = ["Поз.", "Наименование", "Тип", "Кат.", "Pуст\nкВт", "Кс",
               "Pрасч\nкВт", "cosφ", "Кабель", "АВ, А"]
    widths  = [1.2, 5.5, 2.0, 1.0, 1.5, 1.0, 1.5, 1.2, 4.0, 1.5]

    for panel in _iter_panels(vru):
        doc.add_paragraph(f"Щит {panel['id']} — {panel['name']}",
                          style="Heading 2")

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        _set_col_widths(table, widths)
        _add_header_row(table, headers)

        for c in _iter_consumers(panel):
            cable = c.get("cable", {})
            cable_str = ""
            if cable:
                cable_str = (f"{cable.get('mark','')} "
                             f"{cable.get('cores','')}×"
                             f"{cable.get('section_mm2','')}мм²")
            br = c.get("breaker", {})
            br_str = str(br.get("rating", "—")) if br else "—"
            _add_data_row(table, [
                c.get("id", ""),
                c.get("name", ""),
                c.get("type", ""),
                c.get("category_pue", 3),
                f"{c.get('power_kw', 0):.2f}",
                f"{c.get('demand_factor', 1):.2f}",
                f"{c.get('p_calc_kw', 0):.2f}",
                f"{c.get('cos_phi', 1):.2f}",
                cable_str,
                br_str,
            ])

        # Итого по щиту
        _add_data_row(table, [
            "", "ИТОГО:", "", "",
            f"{panel.get('p_installed_kw', 0):.2f}", "",
            f"{panel.get('p_calc_kw', 0):.2f}",
            f"{panel.get('cos_phi', 1):.2f}",
            "", ""
        ], bold=True)

        doc.add_paragraph("")

    # Итого по ВРУ
    doc.add_paragraph(
        f"ИТОГО по ВРУ:  "
        f"Pуст={vru.get('p_installed_kw',0):.2f} кВт  "
        f"Pрасч={vru.get('p_calc_kw',0):.2f} кВт  "
        f"Iвру={vru.get('i_calc_a',0):.2f} А  "
        f"cosφ={vru.get('cos_phi',1):.3f}"
    ).bold = True

    out = docs_dir / "load_table_by_panel.docx"
    doc.save(str(out))
    return out


def _docx_load_summary(project: dict, docs_dir: Path) -> Path:
    vru = project.get("_results", {}).get("vru", {})
    proj_info = project.get("project", {})
    proj_name = proj_info.get("name", "Объект")
    proj_code = proj_info.get("code", "")
    comp = project.get("compensation", {})

    doc = Document()
    title = doc.add_heading("Сводная таблица нагрузок", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"{proj_code} — {proj_name}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    headers = ["Показатель", "Значение"]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    _set_col_widths(table, [7.0, 5.0])
    _add_header_row(table, headers)

    rows_data = [
        ("Pуст, кВт",       f"{vru.get('p_installed_kw', 0):.2f}"),
        ("Pрасч, кВт",      f"{vru.get('p_calc_kw', 0):.2f}"),
        ("Sрасч, кВА",      f"{vru.get('s_calc_kva', 0):.2f}"),
        ("cosφ",            f"{vru.get('cos_phi', 1):.3f}"),
        ("Iвру, А",         f"{vru.get('i_calc_a', 0):.2f}"),
    ]
    br = vru.get("breaker", {})
    if br:
        rows_data.append(("АВ ВРУ", br.get("type", "")))
    cable = vru.get("cable", {})
    if cable:
        rows_data.append(("Кабель ВРУ",
                          f"{cable.get('mark','')} "
                          f"{cable.get('cores','')}×{cable.get('section_mm2','')}мм²"))
    if comp.get("required"):
        krm = comp.get("selected_krm", {})
        rows_data.append(("КРМ", f"{krm.get('model','')} ({krm.get('power_kvar','')} кВАр)"))

    for label, value in rows_data:
        _add_data_row(table, [label, value])

    out = docs_dir / "load_summary.docx"
    doc.save(str(out))
    return out


# ── Публичный API ─────────────────────────────────────────────────────────────

def generate_load_table(project: dict, docs_dir: Path | str,
                         mode: str = "by_panel") -> Path:
    """
    Генерирует ведомость нагрузок.

    Args:
        project:  project.json (с _results)
        docs_dir: папка для сохранения
        mode:     "by_panel" | "by_section" | "summary"

    Returns:
        Path к созданному файлу
    """
    docs_dir = Path(docs_dir)
    docs_dir.mkdir(parents=True, exist_ok=True)

    if mode == "by_section":
        return _txt_load_table_by_section(project, docs_dir)

    if mode == "summary":
        if _DOCX_OK:
            return _docx_load_summary(project, docs_dir)
        return _txt_load_summary(project, docs_dir)

    # by_panel (default)
    if _DOCX_OK:
        return _docx_load_table_by_panel(project, docs_dir)
    return _txt_load_table_by_panel(project, docs_dir)


def generate_all_load_tables(project: dict, docs_dir: Path | str) -> list[Path]:
    """Генерирует все три формата ведомости нагрузок."""
    docs_dir = Path(docs_dir)
    return [
        generate_load_table(project, docs_dir, mode="by_panel"),
        generate_load_table(project, docs_dir, mode="by_section"),
        generate_load_table(project, docs_dir, mode="summary"),
    ]
