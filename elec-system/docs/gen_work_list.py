"""
docs/gen_work_list.py — ведомость объёмов работ по ГОСТ 21.110.

Формат: А4 книжная.
Разделы:
  1. Монтаж щитов и шкафов управления
  2. Прокладка кабелей и проводов (по маркам и типу прокладки)
  3. Разделка и подключение кабелей
  4. Монтаж аппаратуры и оборудования
  5. Испытания и измерения
  6. Прочие работы

Объёмы берутся из _results (длины из length_m_calc, счётчики по data).
"""

from __future__ import annotations

from collections import defaultdict
from pathlib import Path


# ── Сбор объёмов ─────────────────────────────────────────────────────────────

def _build_work_data(project: dict) -> list[dict]:
    """
    Возвращает список позиций ведомости работ.
    Каждая позиция: {"section", "n", "name", "unit", "qty", "note"}
    """
    vru    = project["_results"]["vru"]
    proj   = project.get("project", {})

    # ── Счётчики ──────────────────────────────────────────────────────────────
    panels_all: list[dict]   = []   # все щиты из _results
    n_panels_avr = 0
    cable_agg: dict[tuple, float] = defaultdict(float)  # (mark, install) → length_calc

    def _agg_cable(cb: dict):
        if not cb or not cb.get("section_mm2"):
            return
        mark    = cb.get("mark", "ВВГнг-LS")
        install = cb.get("install", "лоток")
        length  = float(cb.get("length_m_calc") or cb.get("length_m") or 0)
        cable_agg[(mark, install)] += length

    # Вводной кабель ВРУ
    ic = vru.get("incoming_cable", {})
    _agg_cable(ic)
    n_cables_vru = 1 if ic.get("section_mm2") else 0

    # Фидеры → щиты → потребители
    n_panel_cables  = 0
    n_consumer_cables = 0
    n_breakers_total  = 0
    n_contactors      = 0

    for feeder in vru.get("feeders", []):
        for panel_res in feeder.get("panels", []):
            panels_all.append(panel_res)
            pc = panel_res.get("cable", {})
            _agg_cable(pc)
            if pc.get("section_mm2"):
                n_panel_cables += 1

            # Ищем has_avr в исходных данных
            panel_id = panel_res["id"]
            for f_src in project.get("vru", {}).get("feeders", []):
                for p_src in f_src.get("panels", []):
                    if p_src.get("id") == panel_id and p_src.get("has_avr"):
                        n_panels_avr += 1

            for c in panel_res.get("consumers", []):
                cc = c.get("cable", {})
                _agg_cable(cc)
                if cc.get("section_mm2"):
                    n_consumer_cables += 1
                if c.get("breaker", {}).get("rating"):
                    n_breakers_total += 1
                c_type = c.get("type", "")
                if c_type in ("pump", "motor", "hvac", "ventilation_unit", "smoke_fan"):
                    n_contactors += 1

    n_panels = len(panels_all)
    n_cables  = n_cables_vru + n_panel_cables + n_consumer_cables

    # ── Наружные сети ─────────────────────────────────────────────────────────
    outdoor_results = project.get("_results", {}).get("outdoor_networks", [])
    n_outdoor_cables = 0
    for net in outdoor_results:
        cb = net.get("cable", {})
        _agg_cable(cb)
        if cb.get("section_mm2"):
            n_outdoor_cables += 1

    n_cables += n_outdoor_cables

    # ── Суммарные длины лотков (70% кабельных трасс) ─────────────────────────
    total_cable_m = sum(cable_agg.values())
    tray_m   = round(sum(v for (mk, inst), v in cable_agg.items()
                         if "лот" in inst.lower() or "tray" in inst.lower()))
    pipe_m   = round(sum(v for (mk, inst), v in cable_agg.items()
                         if "труб" in inst.lower() or "pipe" in inst.lower()))
    ground_m = round(sum(v for (mk, inst), v in cable_agg.items()
                         if "земл" in inst.lower() or "ground" in inst.lower()))

    # ── Формируем позиции ведомости ───────────────────────────────────────────
    items: list[dict] = []
    _sec: dict[str, list] = defaultdict(list)

    def add(section: str, name: str, unit: str, qty, note: str = ""):
        _sec[section].append({"name": name, "unit": unit, "qty": qty, "note": note})

    # 1. Щиты
    add("1. Монтаж щитов и шкафов управления",
        "Монтаж ВРУ", "шт.", 1,
        f"АВ ввода {vru.get('breaker', {}).get('rating','?')}А")
    add("1. Монтаж щитов и шкафов управления",
        "Монтаж распределительных щитов", "шт.", n_panels,
        "Крепление к стене / на раму")
    if n_panels_avr:
        add("1. Монтаж щитов и шкафов управления",
            "Монтаж блока АВР (ATS)", "шт.", n_panels_avr,
            "Для щитов кат.1")

    # 2. Кабельный лоток
    if tray_m:
        add("2. Прокладка кабелей и проводов",
            "Монтаж кабельного лотка перфорированного", "м",
            round(tray_m * 0.65),  # длина трасс ≈ 65% от суммы кабелей
            "Перфолоток 200×60 мм, с крышкой")
    if pipe_m:
        add("2. Прокладка кабелей и проводов",
            "Прокладка в трубе ПВХ/металлической", "м",
            round(pipe_m * 0.7), "")
    if ground_m:
        add("2. Прокладка кабелей и проводов",
            "Прокладка кабеля в земле (траншея)", "м",
            round(ground_m * 1.1),
            "С засыпкой кирпичом и лентой ПЭ")

    # Прокладка по маркам
    for (mark, install), length_m in sorted(cable_agg.items()):
        add("2. Прокладка кабелей и проводов",
            f"Прокладка {mark}",
            "м",
            round(length_m),
            install)

    # 3. Разделка и подключение
    n_ends = n_cables * 2  # каждый кабель — 2 конца
    add("3. Разделка и подключение кабелей",
        "Разделка кабелей (оконцевание жил)", "конц.", n_ends,
        "Гильзы ГМЛ, наконечники НШВ")
    add("3. Разделка и подключение кабелей",
        "Подключение кабелей к аппаратам и шинам", "конц.", n_ends, "")
    add("3. Разделка и подключение кабелей",
        "Маркировка кабелей и жил", "компл.", 1,
        "Бирки ПВХ, кабельные маркеры")

    # 4. Аппаратура
    add("4. Монтаж аппаратуры и оборудования",
        "Монтаж автоматических выключателей", "шт.",
        n_breakers_total + n_panel_cables + 1,  # потребители + питающие щитов + ВРУ
        "На DIN-рейку 35 мм")
    add("4. Монтаж аппаратуры и оборудования",
        "Монтаж шин N/PE в корпусах щитов", "компл.", n_panels,
        "ШНИ-6×9-12")
    if n_contactors:
        add("4. Монтаж аппаратуры и оборудования",
            "Монтаж контакторов и пускозащитной аппаратуры", "шт.",
            n_contactors, "КМИ, контактор + реле тепловое")
    add("4. Монтаж аппаратуры и оборудования",
        "Заземление корпусов щитов и оборудования (PE)", "точка",
        n_panels + 1, "")

    # 5. Испытания
    add("5. Испытания и измерения",
        "Измерение сопротивления изоляции кабелей", "линия",
        n_cables, "Мегаомметр 1000 В, протокол")
    add("5. Испытания и измерения",
        "Замер полного сопротивления петли фаза-нуль", "точка",
        n_consumer_cables + n_panel_cables, "Прибор типа MI-3143")
    add("5. Испытания и измерения",
        "Проверка работоспособности АВ (расцепление)", "шт.",
        n_breakers_total, "")
    if n_panels_avr:
        add("5. Испытания и измерения",
            "Испытание АВР (имитация пропадания напряжения)", "шт.",
            n_panels_avr, "")

    # 6. Прочие
    add("6. Прочие работы",
        "Визуальный осмотр электроустановки", "компл.", 1,
        "По ПТЭЭП п.3.1")
    add("6. Прочие работы",
        "Сдача исполнительной документации заказчику", "компл.", 1,
        "Комплект ИД: схемы, протоколы, паспорта")

    # ── Разворачиваем в плоский список с нумерацией ───────────────────────────
    result = []
    pos = 1
    for section, section_items in _sec.items():
        result.append({"section": section, "n": None, "name": "",
                        "unit": "", "qty": None, "note": ""})
        for item in section_items:
            result.append({
                "section": None,
                "n":       pos,
                "name":    item["name"],
                "unit":    item["unit"],
                "qty":     item["qty"],
                "note":    item["note"],
            })
            pos += 1

    return result


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _set_cell(cell, text: str, size_pt: int = 10,
              bold: bool = False, center: bool = False):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    run.font.name  = "Times New Roman"


def _add_stamp(doc, proj: dict):
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc.add_paragraph()
    stamp = doc.add_table(rows=5, cols=4)
    stamp.style = "Table Grid"
    col_w   = [4.0, 5.5, 3.0, 2.5]
    hdr_row = ["Должность", "Фамилия И.О.", "Подпись", "Дата"]
    data    = [
        ("Разработал",    proj.get("designer",  "")),
        ("Проверил",      proj.get("checker",   "")),
        ("Нормоконтроль", proj.get("norm_head", "")),
        ("ГИП",           proj.get("gip",       "")),
    ]
    for cell, h, w in zip(stamp.rows[0].cells, hdr_row, col_w):
        cell.width = Cm(w)
        _set_cell(cell, h, bold=True, center=True)
    for i, (role, name) in enumerate(data, start=1):
        row = stamp.rows[i]
        _set_cell(row.cells[0], role)
        _set_cell(row.cells[1], name)


def generate_work_list(project: dict, docs_dir: Path) -> Path:
    """
    Генерирует ведомость объёмов работ.
    Возвращает путь к созданному файлу (.docx или .txt).
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        _has_docx = True
    except ImportError:
        _has_docx = False

    docs_dir = Path(docs_dir)
    docs_dir.mkdir(parents=True, exist_ok=True)

    proj  = project.get("project", {})
    code  = proj.get("code", "ОБЪЕКТ")
    items = _build_work_data(project)

    if not _has_docx:
        txt_path = docs_dir / f"{code}_work_list.txt"
        _write_txt(items, proj, txt_path)
        return txt_path

    out_path = docs_dir / f"{code}_work_list.docx"
    doc = Document()

    # А4 книжная
    sec = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(1.0)
    sec.top_margin    = Cm(1.5)
    sec.bottom_margin = Cm(1.5)

    def _heading(text: str, size: int = 12, bold: bool = False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.name = "Times New Roman"

    _heading("ВЕДОМОСТЬ ОБЪЁМОВ РАБОТ", size=13, bold=True)
    _heading(proj.get("name", ""), size=11)
    _heading(
        f"Код: {code}  |  Стадия: {proj.get('stage','Р')}  |  "
        f"Ред.: {proj.get('revision', 0)}  |  Дата: {proj.get('date','')}",
        size=10,
    )
    doc.add_paragraph()

    # Таблица
    headers  = ["№", "Наименование работы", "Ед.изм.", "Кол-во", "Примечание"]
    col_w_cm = [1.0, 8.5, 2.0, 2.0, 3.5]
    col_w    = [Cm(w) for w in col_w_cm]

    table = doc.add_table(rows=1, cols=5)
    table.style     = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_row = table.rows[0]
    for cell, h, w in zip(hdr_row.cells, headers, col_w):
        cell.width = w
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell(cell, h, bold=True, center=True)

    for item in items:
        row = table.add_row()

        if item["section"]:
            # Строка раздела — объединяем ячейки
            merged = row.cells[0].merge(row.cells[4])
            _set_cell(merged, f"  {item['section']}", bold=True, size_pt=10)
            continue

        qty = item["qty"]
        if isinstance(qty, float) and qty == int(qty):
            qty = int(qty)

        values     = [str(item["n"]), item["name"], item["unit"], str(qty), item["note"]]
        align_flag = [True, False, True, True, False]
        for cell, val, w, center in zip(row.cells, values, col_w, align_flag):
            cell.width = w
            _set_cell(cell, val, center=center)

    _add_stamp(doc, proj)

    doc.save(str(out_path))
    return out_path


# ── TXT fallback ──────────────────────────────────────────────────────────────

def _write_txt(items: list[dict], proj: dict, path: Path):
    code = proj.get("code", "")
    lines = [
        "ВЕДОМОСТЬ ОБЪЁМОВ РАБОТ",
        f"Объект: {proj.get('name','')} | Код: {code}",
        "",
        f"{'№':<4} {'Наименование работы':<50} {'Ед.':<7} {'Кол.':<6} Примечание",
        "-" * 100,
    ]
    for item in items:
        if item["section"]:
            lines.append(f"\n--- {item['section']} ---")
            continue
        qty = item["qty"]
        if isinstance(qty, float) and qty == int(qty):
            qty = int(qty)
        lines.append(
            f"{item['n']:<4} {item['name']:<50} {item['unit']:<7} {qty!s:<6} {item['note']}"
        )
    lines += [
        "",
        f"Разработал: {proj.get('designer','')}",
        f"Проверил:   {proj.get('checker','')}",
    ]
    path.write_text("\n".join(lines), encoding="utf-8")
