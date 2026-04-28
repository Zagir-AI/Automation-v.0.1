"""
docs/gen_pnr.py — программа пусконаладочных работ (ПНР).

Формат: А4 альбомная.
Колонки: № | Наименование работы | Норматив | Ед. | Кол-во | Исполнитель | Срок, дн.

Фазы ПНР:
  1. Подготовительные работы
  2. Визуальный осмотр и проверка монтажа
  3. Измерения и испытания изоляции
  4. Наладка защитной аппаратуры
  5. Комплексные испытания под нагрузкой
  6. Оформление документации

Автоматически добавляются этапы для:
  - АВР (has_avr=True): испытание переключения
  - Кат.1 / FRLS: проверка огнестойкости и категорийности
  - Наружные сети (outdoor_networks): отдельный раздел
"""

from __future__ import annotations

from pathlib import Path
from collections import defaultdict


# ── Сбор данных ───────────────────────────────────────────────────────────────

def _project_stats(project: dict) -> dict:
    """Собирает счётчики для расчёта объёмов ПНР."""
    vru = project["_results"]["vru"]

    n_panels      = 0
    n_consumers   = 0
    n_cables      = 0
    n_breakers    = 0
    n_avr_panels  = 0
    n_cat1        = 0
    has_frls      = False
    has_outdoor   = bool(project.get("_results", {}).get("outdoor_networks"))

    for feeder in vru.get("feeders", []):
        for panel_res in feeder.get("panels", []):
            n_panels += 1

            if panel_res.get("cable", {}).get("section_mm2"):
                n_cables += 1

            # Ищем has_avr и category_pue в исходных данных
            panel_id = panel_res["id"]
            for f_src in project.get("vru", {}).get("feeders", []):
                for p_src in f_src.get("panels", []):
                    if p_src.get("id") == panel_id:
                        if p_src.get("has_avr"):
                            n_avr_panels += 1

            for c in panel_res.get("consumers", []):
                n_consumers += 1
                if c.get("cable", {}).get("section_mm2"):
                    n_cables += 1
                if c.get("breaker", {}).get("rating"):
                    n_breakers += 1
                if c.get("category_pue", 3) == 1:
                    n_cat1 += 1
                mark = (c.get("cable") or {}).get("mark", "")
                if "FRLS" in mark or "FRLSx" in mark:
                    has_frls = True

    # Вводной кабель ВРУ
    ic = vru.get("incoming_cable", {})
    if ic.get("section_mm2"):
        n_cables += 1

    n_measurements = n_cables          # измерений изоляции
    n_loop         = n_consumers + n_panels  # петля ф-0

    return {
        "n_panels":      n_panels,
        "n_consumers":   n_consumers,
        "n_cables":      n_cables,
        "n_breakers":    n_breakers,
        "n_avr_panels":  n_avr_panels,
        "n_cat1":        n_cat1,
        "has_frls":      has_frls,
        "has_outdoor":   has_outdoor,
        "n_measurements": n_measurements,
        "n_loop":        n_loop,
    }


def _build_pnr_stages(project: dict) -> list[dict]:
    """
    Формирует этапы ПНР со структурой разделов.

    Возвращает плоский список: {"section", "n", "name", "norm", "unit", "qty", "exec", "days"}
    Строки с section != None — заголовки разделов.
    """
    st   = _project_stats(project)
    proj = project.get("project", {})

    # Оценка трудоёмкости (дней) на основе объёма
    def _days(n: int, per: int = 10) -> int:
        """Округлённое количество дней: 1 + n // per."""
        return max(1, (n + per - 1) // per)

    _sec: dict[str, list] = defaultdict(list)

    def add(section: str, name: str, norm: str,
            unit: str, qty: int | str, executor: str, days: int | str = 1):
        _sec[section].append({
            "name": name, "norm": norm,
            "unit": unit, "qty": qty,
            "exec": executor, "days": days,
        })

    ORG = "ПНР-организация"
    GEN = "Генподрядчик"
    ZAK = "Заказчик"

    # ── 1. Подготовительные работы ───────────────────────────────────────────
    add("1. Подготовительные работы",
        "Изучение проектной документации и ТУ",
        "СНиП 3.05.06-85 п.4.1", "компл.", 1, ORG, 1)
    add("1. Подготовительные работы",
        "Входной контроль поставленного оборудования",
        "СНиП 3.05.06-85 п.4.2", "позиция",
        st["n_panels"] + 1, GEN, 1)
    add("1. Подготовительные работы",
        "Проверка наличия и комплектности технической документации на оборудование",
        "ПТЭЭП п.1.5.1", "компл.", 1, ORG, 1)

    # ── 2. Визуальный осмотр и проверка монтажа ──────────────────────────────
    add("2. Визуальный осмотр и проверка монтажа",
        "Визуальный осмотр щитов, ВРУ и электрооборудования",
        "ПУЭ п.1.8.1", "компл.", 1, ORG, 1)
    add("2. Визуальный осмотр и проверка монтажа",
        "Проверка соответствия монтажа проектной документации",
        "РД 34.20.501 п.4", "компл.", 1, ORG, 1)
    add("2. Визуальный осмотр и проверка монтажа",
        "Проверка правильности подключения фаз и маркировки шин",
        "ПУЭ п.1.8.20", "точка", st["n_panels"] + 1, ORG, 1)
    add("2. Визуальный осмотр и проверка монтажа",
        "Проверка защитного заземления (PE) всех корпусов",
        "ПУЭ п.1.7.82", "точка", st["n_panels"] + 1, ORG, 1)

    # ── 3. Измерения и испытания изоляции ────────────────────────────────────
    add("3. Измерения и испытания изоляции",
        "Измерение сопротивления изоляции кабелей (Uисп=1000 В)",
        "ПУЭ п.1.8.37", "линия", st["n_measurements"],
        ORG, _days(st["n_measurements"], 15))
    add("3. Измерения и испытания изоляции",
        "Измерение сопротивления петли фаза-нуль",
        "ПУЭ п.1.8.36", "точка", st["n_loop"],
        ORG, _days(st["n_loop"], 20))
    add("3. Измерения и испытания изоляции",
        "Измерение сопротивления заземляющего устройства",
        "ПУЭ п.1.8.39", "замер", 1, ORG, 1)

    if st["has_frls"] or st["n_cat1"]:
        add("3. Измерения и испытания изоляции",
            "Проверка огнестойкости кабельных линий кат.1 (FRLS-кабели)",
            "СП 6.13130.2021 п.4.14", "линия", st["n_cat1"] or 1, ORG, 1)

    # ── 4. Наладка защитной аппаратуры ───────────────────────────────────────
    add("4. Наладка защитной аппаратуры",
        "Проверка срабатывания автоматических выключателей (расцепление)",
        "ГОСТ Р 50345-2010", "шт.", st["n_breakers"], ORG,
        _days(st["n_breakers"], 20))
    add("4. Наладка защитной аппаратуры",
        "Проверка уставок расцепителей и соответствия номиналов АВ",
        "ГОСТ Р 50345-2010", "шт.", st["n_breakers"], ORG, 1)

    if st["n_avr_panels"]:
        add("4. Наладка защитной аппаратуры",
            "Наладка и испытание блоков АВР (имитация пропадания напряжения)",
            "ГОСТ Р МЭК 60947-6-1", "шт.", st["n_avr_panels"], ORG, 1)
        add("4. Наладка защитной аппаратуры",
            "Проверка времени переключения АВР (не более 0.5 с)",
            "ПУЭ п.1.2.18", "шт.", st["n_avr_panels"], ORG, 1)

    # ── 5. Комплексные испытания под нагрузкой ────────────────────────────────
    add("5. Комплексные испытания под нагрузкой",
        "Рабочее включение ВРУ и щитов, проверка токов нагрузки",
        "РД 34.20.501 п.6.2", "компл.", 1, ORG, 1)
    add("5. Комплексные испытания под нагрузкой",
        "Замер рабочих токов и напряжений под нагрузкой",
        "ПТЭЭП п.2.7.7", "точка", st["n_panels"] + 1, ORG, 1)
    add("5. Комплексные испытания под нагрузкой",
        "Проверка симметрии нагрузки по фазам",
        "ПТЭЭП п.2.7.7", "компл.", 1, ORG, 1)

    if st["has_outdoor"]:
        add("5. Комплексные испытания под нагрузкой",
            "Испытание сетей наружного освещения (включение, регулировка)",
            "РД 34.20.501", "компл.", 1, ORG, 1)

    if st["n_avr_panels"]:
        add("5. Комплексные испытания под нагрузкой",
            "Комплексная проверка АВР под нагрузкой",
            "ПУЭ п.1.2.18", "шт.", st["n_avr_panels"], ORG, 1)

    # ── 6. Оформление документации ───────────────────────────────────────────
    add("6. Оформление документации",
        "Оформление протоколов измерений и испытаний",
        "ГОСТ Р 50571.16-2007", "компл.", 1, ORG, 1)
    add("6. Оформление документации",
        "Составление акта технической готовности электроустановки",
        "РД 34.20.501", "акт", 1, ORG + " / " + ZAK, 1)
    add("6. Оформление документации",
        "Передача исполнительной документации заказчику",
        "СНиП 3.05.06-85 п.4.8", "компл.", 1, GEN, 1)
    add("6. Оформление документации",
        "Инструктаж персонала заказчика по эксплуатации ЭУ",
        "ПТЭЭП п.1.4.3", "компл.", 1, ZAK, 1)

    # ── Разворачиваем в плоский список с нумерацией ───────────────────────────
    result = []
    pos    = 1
    for section, items in _sec.items():
        result.append({
            "section": section, "n": None,
            "name": "", "norm": "", "unit": "",
            "qty": None, "exec": "", "days": None,
        })
        for item in items:
            qty = item["qty"]
            if isinstance(qty, float) and qty == int(qty):
                qty = int(qty)
            result.append({
                "section": None,
                "n":       pos,
                "name":    item["name"],
                "norm":    item["norm"],
                "unit":    item["unit"],
                "qty":     qty,
                "exec":    item["exec"],
                "days":    item["days"],
            })
            pos += 1

    return result


# ── DOCX ──────────────────────────────────────────────────────────────────────

_HEADERS  = ["№", "Наименование работы", "Норматив", "Ед.", "Кол.", "Исполнитель", "Срок, дн."]
_COL_W_CM = [0.8, 9.5, 3.8, 1.5, 1.5, 4.0, 1.8]


def _set_cell(cell, text: str, size_pt: int = 10,
              bold: bool = False, center: bool = False):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    run.font.name  = "Times New Roman"


def _add_stamp(doc, proj: dict):
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc.add_paragraph()
    stamp   = doc.add_table(rows=5, cols=4)
    stamp.style = "Table Grid"
    col_w   = [4.0, 5.5, 3.0, 2.5]
    hdr_row = ["Должность", "Фамилия И.О.", "Подпись", "Дата"]
    data    = [
        ("Разработал",    proj.get("designer",  "")),
        ("Проверил",      proj.get("checker",   "")),
        ("Нормоконтроль", proj.get("norm_head", "")),
        ("ГИП",           proj.get("gip",       "")),
    ]
    for cell, h, w in zip(stamp.rows[0].cells, hdr_row, col_w):
        cell.width = Cm(w)
        _set_cell(cell, h, bold=True, center=True)
    for i, (role, name) in enumerate(data, start=1):
        row = stamp.rows[i]
        _set_cell(row.cells[0], role)
        _set_cell(row.cells[1], name)


def generate_pnr(project: dict, docs_dir: Path) -> Path:
    """
    Генерирует программу ПНР.
    Возвращает путь к созданному файлу (.docx или .txt).
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        _has_docx = True
    except ImportError:
        _has_docx = False

    docs_dir = Path(docs_dir)
    docs_dir.mkdir(parents=True, exist_ok=True)

    proj   = project.get("project", {})
    code   = proj.get("code", "ОБЪЕКТ")
    stages = _build_pnr_stages(project)

    if not _has_docx:
        txt_path = docs_dir / f"{code}_pnr.txt"
        _write_txt(stages, proj, txt_path)
        return txt_path

    out_path = docs_dir / f"{code}_pnr.docx"
    doc = Document()

    # А4 альбомная
    sec = doc.sections[0]
    sec.page_width    = Cm(29.7)
    sec.page_height   = Cm(21.0)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(1.0)
    sec.top_margin    = Cm(1.5)
    sec.bottom_margin = Cm(1.5)

    def _heading(text: str, size: int = 12, bold: bool = False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.name = "Times New Roman"

    _heading("ПРОГРАММА ПУСКОНАЛАДОЧНЫХ РАБОТ", size=13, bold=True)
    _heading(proj.get("name", ""), size=11)
    _heading(
        f"Код: {code}  |  Стадия: {proj.get('stage','Р')}  |  "
        f"Ред.: {proj.get('revision', 0)}  |  Дата: {proj.get('date','')}",
        size=10,
    )
    doc.add_paragraph()

    # Таблица
    col_w = [Cm(w) for w in _COL_W_CM]
    table = doc.add_table(rows=1, cols=len(_HEADERS))
    table.style     = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_row = table.rows[0]
    for cell, h, w in zip(hdr_row.cells, _HEADERS, col_w):
        cell.width = w
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell(cell, h, bold=True, center=True)

    for item in stages:
        row = table.add_row()

        if item["section"]:
            merged = row.cells[0].merge(row.cells[len(_HEADERS) - 1])
            _set_cell(merged, f"  {item['section']}", bold=True)
            continue

        values     = [str(item["n"]), item["name"], item["norm"],
                      item["unit"], str(item["qty"]), item["exec"], str(item["days"])]
        align_flag = [True, False, False, True, True, False, True]

        for cell, val, w, center in zip(row.cells, values, col_w, align_flag):
            cell.width = w
            _set_cell(cell, val, center=center)

    # Итоговая строка: суммарные дни
    total_days = sum(item["days"] for item in stages if item["days"] is not None)
    doc.add_paragraph()
    p_total = doc.add_paragraph()
    p_total.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rt = p_total.add_run(
        f"Итого этапов ПНР: {sum(1 for i in stages if i['n'] is not None)}  |  "
        f"Расчётная продолжительность: {total_days} рабочих дней"
    )
    rt.font.size = Pt(10)
    rt.font.name = "Times New Roman"

    _add_stamp(doc, proj)

    doc.save(str(out_path))
    return out_path


# ── TXT fallback ──────────────────────────────────────────────────────────────

def _write_txt(stages: list[dict], proj: dict, path: Path):
    code = proj.get("code", "")
    lines = [
        "ПРОГРАММА ПУСКОНАЛАДОЧНЫХ РАБОТ",
        f"Объект: {proj.get('name','')} | Код: {code}",
        "",
        f"{'№':<4} {'Наименование':<50} {'Норматив':<18} {'Ед.':<7} {'Кол.':<5} {'Исполнитель':<20} Дней",
        "-" * 120,
    ]
    for item in stages:
        if item["section"]:
            lines.append(f"\n--- {item['section']} ---")
            continue
        lines.append(
            f"{item['n']:<4} {item['name']:<50} {item['norm']:<18} "
            f"{item['unit']:<7} {item['qty']!s:<5} {item['exec']:<20} {item['days']}"
        )
    total_days = sum(item["days"] for item in stages if item["days"] is not None)
    lines += [
        "",
        f"Расчётная продолжительность ПНР: {total_days} рабочих дней",
        "",
        f"Разработал: {proj.get('designer','')}",
        f"Проверил:   {proj.get('checker','')}",
    ]
    path.write_text("\n".join(lines), encoding="utf-8")
