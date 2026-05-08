"""
docs/gen_spec.py — спецификация оборудования, изделий и материалов.

Формат: ГОСТ 21.110-2013, форма 3 (А4, альбомная).
Разделы:
  1. Щиты и шкафы управления
  2. Аппараты защиты
  3. Кабели и провода
  4. Электроустановочные изделия и оборудование
  5. Прочие материалы

Длины кабелей берутся из cable["length_m_calc"] (уже с запасом и вертикальными участками).
Обозначения АВ — из get_breaker_designation() по серии производителя.
Шаблонные позиции щита — из get_template_items() по типу щита.
"""

from __future__ import annotations

import sys
from pathlib import Path
from collections import defaultdict

# ── Вспомогательные импорты ───────────────────────────────────────────────────

def _import_breaker_tables():
    here = Path(__file__).resolve().parent.parent
    if str(here) not in sys.path:
        sys.path.insert(0, str(here))
    try:
        from data.breakers.breaker_tables import get_breaker_designation
        return get_breaker_designation
    except ImportError:
        return None


def _import_spec_templates():
    here = Path(__file__).resolve().parent.parent
    if str(here) not in sys.path:
        sys.path.insert(0, str(here))
    try:
        from data.spec_templates import get_template_items
        return get_template_items
    except ImportError:
        return None


# ── Эффективная длина кабеля ──────────────────────────────────────────────────

def _effective_length(cable: dict) -> float:
    """Возвращает расчётную длину кабеля (уже с запасом если есть)."""
    calc = cable.get("length_m_calc")
    if calc is not None:
        return float(calc)
    return float(cable.get("length_m", 0))


# ── Сбор данных спецификации ──────────────────────────────────────────────────

def _build_spec_data(project: dict) -> dict:
    """
    Обходит _results и собирает все позиции спецификации.

    Returns:
        {
          "panels":   [{"id", "name", "note", "has_avr"}],
          "breakers": {(rating, char, poles, series): {"count", "mark", "name", "gost"}},
          "cables":   {(mark, cores, section): {"length_m", "routing_notes": set}},
          "hardware": [{"name", "mark", "unit", "qty", "note", "section"}],
          "extra":    list from project["extra_items"],
          "series":   str brand,
        }
    """
    get_breaker_designation = _import_breaker_tables()
    get_template_items      = _import_spec_templates()

    vru_result = project.get("_results", {}).get("vru", {})
    proj       = project.get("project", {})
    series_brand = proj.get("breaker_series", "IEK")

    panels   = []
    breakers: dict[tuple, dict] = {}
    cables:   dict[tuple, dict] = {}
    hardware: list[dict]        = []

    # ── helpers ──────────────────────────────────────────────────────────────

    def _br_key(br: dict) -> tuple:
        return (br["rating"], br.get("char", "C"), br.get("poles", 3), series_brand)

    def _add_breaker(br: dict, count: int = 1):
        if not br or not br.get("rating"):
            return
        key = _br_key(br)
        if key not in breakers:
            if get_breaker_designation:
                desig = get_breaker_designation(br["rating"], br.get("char", "C"),
                                                br.get("poles", 3), series_brand)
            else:
                desig = {
                    "mark": f"АВ {br['rating']}А {br.get('char','C')}",
                    "name": (f"Выключатель автоматический {br.get('poles',3)}П "
                             f"{br['rating']}А хар.{br.get('char','C')}"),
                    "gost": "ГОСТ IEC 60898-1",
                }
            breakers[key] = {"count": 0, **desig}
        breakers[key]["count"] += count

    def _add_cable(cb: dict):
        if not cb or not cb.get("section_mm2"):
            return
        mark    = cb.get("mark", "ВВГнг-LS")
        cores   = cb.get("cores", 4)
        section = cb.get("section_mm2")
        length  = _effective_length(cb)
        note    = cb.get("routing_note", "")
        key = (mark, cores, section)
        if key not in cables:
            cables[key] = {"length_m": 0.0, "routing_notes": set()}
        cables[key]["length_m"]    += length
        if note:
            cables[key]["routing_notes"].add(note)

    # ── ВРУ ──────────────────────────────────────────────────────────────────

    vru_src = project.get("vru", {})
    vru_br  = vru_result.get("breaker", {})
    vru_cb  = vru_result.get("incoming_cable", {})

    panels.append({
        "id":      vru_result.get("id", vru_src.get("id", "ВРУ-1")),
        "name":    vru_src.get("name", "Вводно-распределительное устройство"),
        "note":    (f"Iн={vru_br.get('rating','?')}А, "
                    f"Iкз={vru_src.get('isc_ka', 10)}кА"),
        "has_avr": vru_src.get("has_avr", False),
    })
    _add_breaker(vru_br)
    _add_cable(vru_cb)

    # ── Фидеры → Щиты → Потребители ──────────────────────────────────────────

    for feeder in vru_result.get("feeders", []):
        for panel_res in feeder.get("panels", []):
            panel_id   = panel_res.get("id", "")
            panel_name = panel_res.get("name", panel_id)
            panel_src  = {}

            # Найдём источник щита в project для panel_type
            for f_src in project.get("vru", {}).get("feeders", []):
                for p_src in f_src.get("panels", []):
                    if p_src.get("id") == panel_id:
                        panel_src = p_src
                        break

            panel_type  = panel_src.get("panel_type") or panel_src.get("type", "")
            has_avr     = panel_src.get("has_avr", False)
            bus_a       = panel_res.get("breaker", {}).get("rating", "?")
            consumers   = panel_res.get("consumers", [])
            n_consumers = sum(1 for c in consumers if not c.get("reserve"))

            panels.append({
                "id":      panel_id,
                "name":    panel_name,
                "note":    f"Iн шины={bus_a}А",
                "has_avr": has_avr,
            })

            # АВ питания щита
            _add_breaker(panel_res.get("breaker", {}))

            # Кабель питания щита
            _add_cable(panel_res.get("cable", {}))

            # Потребители
            n_breakers = 0
            for c in consumers:
                br = c.get("breaker", {})
                if br and br.get("rating"):
                    _add_breaker(br)
                    n_breakers += 1
                _add_cable(c.get("cable", {}))

            # Шаблонные позиции по типу щита (always apply "all", even if no specific type)
            if get_template_items:
                tmpl_items = get_template_items(panel_type, n_consumers, n_breakers)
                for ti in tmpl_items:
                    hardware.append(ti)

    # ── Наружные сети освещения ───────────────────────────────────────────────

    outdoor_results = project.get("_results", {}).get("outdoor_networks", [])
    for net in outdoor_results:
        _add_cable(net.get("cable", {}))

    # ── extra_items из project.json ───────────────────────────────────────────

    extra = project.get("extra_items", [])

    return {
        "panels":   panels,
        "breakers": breakers,
        "cables":   cables,
        "hardware": hardware,
        "extra":    extra,
        "series":   series_brand,
    }


# ── DOCX ──────────────────────────────────────────────────────────────────────

_COL_WIDTHS_CM = [1.5, 4.0, 8.5, 1.5, 1.2, 5.3]   # Поз. Марка Наим. Кол. Ед. Примеч.


def _cell_text(cell, text: str, size_pt: int = 10, bold: bool = False,
               center: bool = False):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    run.font.name  = "Times New Roman"


def _add_header_row(table, col_widths):
    from docx.shared import Cm
    from docx.enum.table import WD_ALIGN_VERTICAL
    headers = ["Поз.", "Марка / Обозначение", "Наименование", "Кол.", "Ед.", "Примечание"]
    row = table.rows[0]
    for cell, h, w in zip(row.cells, headers, col_widths):
        cell.width = Cm(w)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _cell_text(cell, h, bold=True, center=True)


def _add_section_row(table, title: str):
    row = table.add_row()
    merged = row.cells[0].merge(row.cells[5])
    _cell_text(merged, f"  {title}", bold=True)


def _add_item_row(table, pos: int, mark: str, name: str,
                  qty, unit: str, note: str, col_widths):
    from docx.shared import Cm
    row = table.add_row()
    data = [str(pos), mark, name, str(qty), unit, note]
    for cell, val, w in zip(row.cells, data, col_widths):
        cell.width = Cm(w)
        center = val in (str(pos), str(qty), unit)
        _cell_text(cell, val, center=center)
    return pos + 1


def _add_stamp(doc, proj: dict):
    """
    Штамп (подписи): таблица 4 строки.
    Колонки: Должность | ФИО | Подпись | Дата
    """
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc.add_paragraph()
    stamp = doc.add_table(rows=5, cols=4)
    stamp.style = "Table Grid"

    col_w = [4.0, 5.5, 3.0, 2.5]
    headers_r0 = ["Должность", "Фамилия И.О.", "Подпись", "Дата"]
    rows_data = [
        ("Разработал",    proj.get("designer",   "")),
        ("Проверил",      proj.get("checker",    "")),
        ("Нормоконтроль", proj.get("norm_head",  "")),
        ("ГИП",           proj.get("gip",        "")),
    ]

    def _st_cell(cell, text, bold=False, center=False):
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        run.font.bold = bold
        run.font.name = "Times New Roman"

    # Строка заголовков
    for cell, h, w in zip(stamp.rows[0].cells, headers_r0, col_w):
        from docx.shared import Cm
        cell.width = Cm(w)
        _st_cell(cell, h, bold=True, center=True)

    for i, (role, name) in enumerate(rows_data, start=1):
        row = stamp.rows[i]
        _st_cell(row.cells[0], role)
        _st_cell(row.cells[1], name)
        # Колонки "Подпись" и "Дата" оставляем пустыми (заполняются вручную)

    # Строка: организация + дата
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_org.add_run(
        f"Организация: {proj.get('org','')}   |   "
        f"Город: {proj.get('city','')}   |   "
        f"Стадия: {proj.get('stage','Р')}   |   "
        f"Объект: {proj.get('object_type','')}   |   "
        f"Система: {proj.get('system','ЭС и ЭО')}"
    )
    run.font.size = Pt(9)
    run.font.name = "Times New Roman"


# ── Публичный API ─────────────────────────────────────────────────────────────

def generate_spec(project: dict, docs_dir: Path) -> Path:
    """
    Генерирует спецификацию оборудования по ГОСТ 21.110-2013.
    Возвращает путь к созданному файлу (.docx или .txt).
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        _has_docx = True
    except ImportError:
        _has_docx = False

    docs_dir = Path(docs_dir)
    docs_dir.mkdir(parents=True, exist_ok=True)

    proj  = project.get("project", {})
    code  = proj.get("code", "ОБЪЕКТ")
    name  = proj.get("name", "")
    stage = proj.get("stage", "Р")
    rev   = proj.get("revision", 0)
    date  = proj.get("date", "")

    data = _build_spec_data(project)

    if not _has_docx:
        txt_path = docs_dir / f"{code}_spec.txt"
        _write_txt(data, proj, txt_path)
        return txt_path

    out_path = docs_dir / f"{code}_spec.docx"
    doc = Document()

    # ── Страница: А4 альбомная ────────────────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width    = Cm(29.7)
    sec.page_height   = Cm(21.0)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(1.0)
    sec.top_margin    = Cm(1.5)
    sec.bottom_margin = Cm(1.5)

    def _heading(text: str, size: int = 12, bold: bool = False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.name = "Times New Roman"

    _heading("СПЕЦИФИКАЦИЯ ОБОРУДОВАНИЯ, ИЗДЕЛИЙ И МАТЕРИАЛОВ", size=13, bold=True)
    _heading(name, size=11)
    _heading(
        f"Раздел: {proj.get('system','ЭС и ЭО')}  |  Стадия: {stage}  |  "
        f"Код: {code}  |  Ред.: {rev}  |  Дата: {date}",
        size=10,
    )
    doc.add_paragraph()

    # ── Основная таблица ──────────────────────────────────────────────────────
    col_widths = _COL_WIDTHS_CM
    table = doc.add_table(rows=1, cols=6)
    table.style     = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_header_row(table, col_widths)

    pos = 1

    # ── 1. Щиты и шкафы ──────────────────────────────────────────────────────
    _add_section_row(table, "1. Щиты и шкафы управления")
    for panel in data["panels"]:
        avr_note = " (с АВР)" if panel.get("has_avr") else ""
        pos = _add_item_row(table, pos,
                             panel["id"],
                             panel["name"] + avr_note,
                             1, "шт.",
                             panel.get("note", ""),
                             col_widths)

    # ── 2. Аппараты защиты ───────────────────────────────────────────────────
    _add_section_row(table, "2. Аппараты защиты")
    for key in sorted(data["breakers"]):
        br   = data["breakers"][key]
        gost = br.get("gost", "")
        pos  = _add_item_row(table, pos,
                              br["mark"],
                              br["name"],
                              br["count"], "шт.",
                              gost,
                              col_widths)

    # ── 3. Кабели и провода ───────────────────────────────────────────────────
    _add_section_row(table, "3. Кабели и провода")
    for (mark, cores, section) in sorted(data["cables"]):
        entry  = data["cables"][(mark, cores, section)]
        length = round(entry["length_m"])
        notes  = "; ".join(sorted(entry["routing_notes"]))
        cable_name = f"Кабель {mark} {cores}×{section} мм²"
        cable_mark = f"{mark} {cores}×{section}"
        pos = _add_item_row(table, pos,
                             cable_mark,
                             cable_name,
                             length, "м",
                             notes or "с запасом 20%",
                             col_widths)

    # ── 4. Электроустановочные изделия (шаблоны щитов) ───────────────────────
    if data["hardware"]:
        _add_section_row(table, "4. Электроустановочные изделия и оборудование")
        # Группируем одинаковые позиции
        hw_agg: dict[tuple, dict] = {}
        for item in data["hardware"]:
            key = (item["name"], item.get("mark", ""), item["unit"])
            if key not in hw_agg:
                hw_agg[key] = {**item, "qty": 0}
            hw_agg[key]["qty"] += item["qty"]
        for key in sorted(hw_agg):
            item = hw_agg[key]
            qty  = item["qty"]
            if isinstance(qty, float) and qty == int(qty):
                qty = int(qty)
            pos = _add_item_row(table, pos,
                                 item.get("mark", ""),
                                 item["name"],
                                 qty, item["unit"],
                                 item.get("note", ""),
                                 col_widths)

    # ── 5. Прочие материалы (extra_items) ────────────────────────────────────
    if data["extra"]:
        _add_section_row(table, "5. Прочие материалы")
        for item in data["extra"]:
            qty = item.get("qty", 1)
            if isinstance(qty, float) and qty == int(qty):
                qty = int(qty)
            pos = _add_item_row(table, pos,
                                 item.get("mark", ""),
                                 item.get("name", ""),
                                 qty,
                                 item.get("unit", "шт."),
                                 item.get("note", ""),
                                 col_widths)

    doc.add_paragraph()
    _add_stamp(doc, proj)

    doc.save(str(out_path))
    return out_path


# ── XLSX ──────────────────────────────────────────────────────────────────────

def generate_spec_xlsx(project: dict, docs_dir: Path) -> Path:
    """
    Генерирует спецификацию оборудования в формате xlsx (ГОСТ 21.110-2013).
    Рабочий документ для комплектации; штамп не включается.
    Возвращает путь к созданному файлу.
    """
    from openpyxl import Workbook
    from openpyxl.styles import (Font, PatternFill, Alignment, Border, Side,
                                  GradientFill)
    from openpyxl.utils import get_column_letter

    docs_dir = Path(docs_dir)
    docs_dir.mkdir(parents=True, exist_ok=True)

    proj  = project.get("project", {})
    code  = proj.get("code", "ОБЪЕКТ")
    name  = proj.get("name", "")
    stage = proj.get("stage", "Р")
    rev   = proj.get("revision", 0)
    date  = proj.get("date", "")

    data = _build_spec_data(project)

    out_path = docs_dir / f"{code}_spec.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Спецификация"

    # ── Стили ────────────────────────────────────────────────────────────────
    TNR = "Times New Roman"
    thin = Side(style="thin")
    border_all = Border(left=thin, right=thin, top=thin, bottom=thin)

    def _font(size=10, bold=False, italic=False, color="FF000000"):
        return Font(name=TNR, size=size, bold=bold, italic=italic, color=color)

    def _fill(hex_color):
        return PatternFill("solid", fgColor=hex_color)

    def _align(horizontal="left", wrap=True):
        return Alignment(horizontal=horizontal, vertical="center",
                         wrap_text=wrap)

    fill_header  = _fill("FFCCCCCC")
    fill_section = _fill("FFD6E4F7")

    # ── Ширины колонок (символы) ──────────────────────────────────────────────
    col_widths_ch = [6, 20, 45, 8, 6, 30]
    col_letters   = [get_column_letter(i + 1) for i in range(6)]
    for letter, w in zip(col_letters, col_widths_ch):
        ws.column_dimensions[letter].width = w

    # ── Шапка документа (строки 1–4) ─────────────────────────────────────────
    def _header_row(row_num: int, text: str, size: int = 10, bold: bool = False):
        ws.merge_cells(f"A{row_num}:F{row_num}")
        cell = ws.cell(row=row_num, column=1, value=text)
        cell.font      = _font(size=size, bold=bold)
        cell.alignment = _align(horizontal="center")
        ws.row_dimensions[row_num].height = 18

    _header_row(1, "СПЕЦИФИКАЦИЯ ОБОРУДОВАНИЯ, ИЗДЕЛИЙ И МАТЕРИАЛОВ",
                size=13, bold=True)
    _header_row(2, name, size=11)
    _header_row(3,
        f"Код: {code}  |  Стадия: {stage}  |  Ред.: {rev}  |  Дата: {date}",
        size=10)
    ws.row_dimensions[4].height = 8  # пустой отступ

    # ── Строка заголовков таблицы (row 5) ────────────────────────────────────
    headers = ["Поз.", "Марка / Обозначение", "Наименование",
               "Кол.", "Ед.", "Примечание"]
    for col_idx, h in enumerate(headers, start=1):
        cell = ws.cell(row=5, column=col_idx, value=h)
        cell.font      = _font(size=10, bold=True)
        cell.fill      = fill_header
        cell.alignment = _align(horizontal="center")
        cell.border    = border_all
    ws.row_dimensions[5].height = 25

    # freeze: 4 строки шапки + 1 строка заголовков → данные с 6-й строки
    ws.freeze_panes = "A6"

    # ── Заполнение данных ─────────────────────────────────────────────────────
    current_row = 6
    pos = 1

    def _section_row(title: str):
        nonlocal current_row
        ws.merge_cells(f"A{current_row}:F{current_row}")
        cell = ws.cell(row=current_row, column=1, value=f"  {title}")
        cell.font      = _font(size=10, bold=True)
        cell.fill      = fill_section
        cell.alignment = _align(horizontal="left")
        cell.border    = border_all
        ws.row_dimensions[current_row].height = 18
        current_row += 1

    def _data_row(p: int, mark: str, name_s: str, qty, unit: str, note: str):
        nonlocal current_row, pos
        values = [str(p), mark, name_s, str(qty), unit, note]
        aligns = ["center", "left", "left", "center", "center", "left"]
        for col_idx, (val, aln) in enumerate(zip(values, aligns), start=1):
            cell = ws.cell(row=current_row, column=col_idx, value=val)
            cell.font      = _font(size=10)
            cell.alignment = _align(horizontal=aln)
            cell.border    = border_all
        ws.row_dimensions[current_row].height = 16
        current_row += 1
        pos += 1

    # 1. Щиты и шкафы
    _section_row("1. Щиты и шкафы управления")
    for panel in data["panels"]:
        avr_note = " (с АВР)" if panel.get("has_avr") else ""
        _data_row(pos, panel["id"], panel["name"] + avr_note,
                  1, "шт.", panel.get("note", ""))

    # 2. Аппараты защиты
    _section_row("2. Аппараты защиты")
    for key in sorted(data["breakers"]):
        br = data["breakers"][key]
        _data_row(pos, br["mark"], br["name"],
                  br["count"], "шт.", br.get("gost", ""))

    # 3. Кабели и провода
    _section_row("3. Кабели и провода")
    for (mark, cores, section) in sorted(data["cables"]):
        entry  = data["cables"][(mark, cores, section)]
        length = round(entry["length_m"])
        notes  = "; ".join(sorted(entry["routing_notes"]))
        cable_mark = f"{mark} {cores}×{section}"
        cable_name = f"Кабель {mark} {cores}×{section} мм²"
        _data_row(pos, cable_mark, cable_name,
                  length, "м", notes or "с запасом 20%")

    # 4. Электроустановочные изделия
    if data["hardware"]:
        _section_row("4. Электроустановочные изделия и оборудование")
        hw_agg: dict[tuple, dict] = {}
        for item in data["hardware"]:
            key = (item["name"], item.get("mark", ""), item["unit"])
            if key not in hw_agg:
                hw_agg[key] = {**item, "qty": 0}
            hw_agg[key]["qty"] += item["qty"]
        for key in sorted(hw_agg):
            item = hw_agg[key]
            qty  = item["qty"]
            if isinstance(qty, float) and qty == int(qty):
                qty = int(qty)
            _data_row(pos, item.get("mark", ""), item["name"],
                      qty, item["unit"], item.get("note", ""))

    # 5. Прочие материалы
    if data["extra"]:
        _section_row("5. Прочие материалы")
        for item in data["extra"]:
            qty = item.get("qty", 1)
            if isinstance(qty, float) and qty == int(qty):
                qty = int(qty)
            _data_row(pos, item.get("mark", ""), item.get("name", ""),
                      qty, item.get("unit", "шт."), item.get("note", ""))

    wb.save(str(out_path))
    return out_path


# ── TXT fallback ──────────────────────────────────────────────────────────────

def _write_txt(data: dict, proj: dict, path: Path):
    """Запасной вариант без python-docx."""
    code  = proj.get("code", "")
    name  = proj.get("name", "")
    stage = proj.get("stage", "Р")
    rev   = proj.get("revision", 0)

    lines = [
        "СПЕЦИФИКАЦИЯ ОБОРУДОВАНИЯ, ИЗДЕЛИЙ И МАТЕРИАЛОВ",
        f"Объект: {name}",
        f"Код: {code}  |  Стадия: {stage}  |  Ред.: {rev}",
        "",
        f"{'Поз.':<5} {'Марка':<22} {'Наименование':<45} {'Кол.':<7} {'Ед.':<5} Примечание",
        "-" * 110,
    ]
    pos = 1

    def row(mark, name_s, qty, unit, note=""):
        nonlocal pos
        lines.append(f"{pos:<5} {mark:<22} {name_s:<45} {qty!s:<7} {unit:<5} {note}")
        pos += 1

    lines.append("--- 1. Щиты и шкафы управления ---")
    for panel in data["panels"]:
        avr = " (АВР)" if panel.get("has_avr") else ""
        row(panel["id"], panel["name"] + avr, 1, "шт.", panel.get("note", ""))

    lines.append("--- 2. Аппараты защиты ---")
    for key in sorted(data["breakers"]):
        br = data["breakers"][key]
        row(br["mark"], br["name"], br["count"], "шт.", br.get("gost", ""))

    lines.append("--- 3. Кабели и провода ---")
    for (mark, cores, section) in sorted(data["cables"]):
        entry  = data["cables"][(mark, cores, section)]
        length = round(entry["length_m"])
        notes  = "; ".join(sorted(entry["routing_notes"]))
        cable_mark = f"{mark} {cores}×{section}"
        cable_name = f"Кабель {mark} {cores}×{section} мм²"
        row(cable_mark, cable_name, length, "м", notes or "с запасом 20%")

    if data["hardware"]:
        lines.append("--- 4. Электроустановочные изделия ---")
        hw_agg: dict[tuple, dict] = {}
        for item in data["hardware"]:
            key = (item["name"], item.get("mark", ""), item["unit"])
            if key not in hw_agg:
                hw_agg[key] = {**item, "qty": 0}
            hw_agg[key]["qty"] += item["qty"]
        for key in sorted(hw_agg):
            item = hw_agg[key]
            qty  = item["qty"]
            if isinstance(qty, float) and qty == int(qty):
                qty = int(qty)
            row(item.get("mark", ""), item["name"], qty, item["unit"], item.get("note", ""))

    if data["extra"]:
        lines.append("--- 5. Прочие материалы ---")
        for item in data["extra"]:
            qty = item.get("qty", 1)
            if isinstance(qty, float) and qty == int(qty):
                qty = int(qty)
            row(item.get("mark", ""), item.get("name", ""),
                qty, item.get("unit", "шт."), item.get("note", ""))

    lines += [
        "",
        f"Разработал: {proj.get('designer','')}",
        f"Проверил:   {proj.get('checker','')}",
        f"Н.контроль: {proj.get('norm_head','')}",
        f"ГИП:        {proj.get('gip','')}",
    ]
    path.write_text("\n".join(lines), encoding="utf-8")
