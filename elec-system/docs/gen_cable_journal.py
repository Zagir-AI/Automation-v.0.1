"""
docs/gen_cable_journal.py — кабельный журнал по ГОСТ 21.613.

Формат: А3 альбомная.
Колонки: № | № КЛ | Марка | Жилы×Сечение | Откуда | Куда | L план,м | L расч.,м |
         Прокладка | АВ,А | ΔU,% | Примечание

Группировка по разделам (ЭОМ / ОВ / ВК / ...) и щитам.
Длины: length_m_plan — по чертежу, length_m_calc — с запасом/стояком.
Номера КЛ из project["cable_numbering"] (если есть).
"""

from __future__ import annotations

from pathlib import Path


# ── Сбор данных ───────────────────────────────────────────────────────────────

def _collect_cables(project: dict) -> list[dict]:
    """
    Собирает все кабельные линии из _results в плоский список.

    Каждая запись:
      n, cable_no, mark, cores, section, from_id, from_name,
      to_id, to_name, length_plan, length_calc, install,
      breaker_a, voltage_drop, section_tag, note, is_reserve
    """
    vru        = project["_results"]["vru"]
    numbering  = project.get("cable_numbering", {})
    records    = []
    seq        = 1

    def _cable_no(consumer_id: str | None) -> str:
        if consumer_id and consumer_id in numbering:
            return numbering[consumer_id]
        return "—"

    def _len_plan(cb: dict) -> float:
        return float(cb.get("length_m_plan") or cb.get("length_m") or 0)

    def _len_calc(cb: dict) -> float:
        calc = cb.get("length_m_calc")
        if calc is not None:
            return float(calc)
        return float(cb.get("length_m") or 0)

    # ── Вводной кабель ВРУ ───────────────────────────────────────────────────
    ic = vru.get("incoming_cable", {})
    if ic.get("section_mm2"):
        vru_br = vru.get("breaker", {})
        records.append({
            "n":            seq,
            "cable_no":     "—",
            "mark":         ic.get("mark", "ВВГнг-LS"),
            "cores":        ic.get("cores", 4),
            "section":      ic.get("section_mm2"),
            "from_id":      "ТП-1",
            "from_name":    "Трансформаторная подстанция",
            "to_id":        vru["id"],
            "to_name":      vru.get("name", vru["id"]),
            "length_plan":  _len_plan(ic),
            "length_calc":  _len_calc(ic),
            "install":      ic.get("install", "лоток"),
            "breaker_a":    vru_br.get("rating", "—"),
            "voltage_drop": ic.get("voltage_drop_pct", "—"),
            "section_tag":  "ЭОМ",
            "note":         ic.get("routing_note", ""),
            "is_reserve":   False,
        })
        seq += 1

    # ── Фидеры → Щиты → Потребители ─────────────────────────────────────────
    for feeder in vru.get("feeders", []):
        f_section = feeder.get("section", "ЭОМ")

        for panel_res in feeder.get("panels", []):
            panel_id   = panel_res["id"]
            panel_name = panel_res.get("name", panel_id)
            panel_br   = panel_res.get("breaker", {})

            # Питающий кабель щита
            pc = panel_res.get("cable", {})
            if pc.get("section_mm2"):
                records.append({
                    "n":            seq,
                    "cable_no":     "—",
                    "mark":         pc.get("mark", "ВВГнг-LS"),
                    "cores":        pc.get("cores", 4),
                    "section":      pc.get("section_mm2"),
                    "from_id":      vru["id"],
                    "from_name":    vru.get("name", vru["id"]),
                    "to_id":        panel_id,
                    "to_name":      panel_name,
                    "length_plan":  _len_plan(pc),
                    "length_calc":  _len_calc(pc),
                    "install":      pc.get("install", "лоток"),
                    "breaker_a":    panel_br.get("rating", "—"),
                    "voltage_drop": pc.get("voltage_drop_pct", "—"),
                    "section_tag":  f_section,
                    "note":         pc.get("routing_note", ""),
                    "is_reserve":   False,
                })
                seq += 1

            # Кабели потребителей
            for c in panel_res.get("consumers", []):
                cc = c.get("cable", {})
                if not cc.get("section_mm2"):
                    continue

                c_br  = c.get("breaker", {})
                c_sec = c.get("section", f_section) or f_section
                note  = cc.get("routing_note", "")
                if c.get("reserve"):
                    note = ("резерв; " + note).strip("; ")

                records.append({
                    "n":            seq,
                    "cable_no":     _cable_no(c["id"]),
                    "mark":         cc.get("mark", "ВВГнг-LS"),
                    "cores":        cc.get("cores", 3),
                    "section":      cc.get("section_mm2"),
                    "from_id":      panel_id,
                    "from_name":    panel_name,
                    "to_id":        c["id"],
                    "to_name":      c.get("name", c["id"]),
                    "length_plan":  _len_plan(cc),
                    "length_calc":  _len_calc(cc),
                    "install":      cc.get("install", "лоток"),
                    "breaker_a":    c_br.get("rating", "—"),
                    "voltage_drop": cc.get("voltage_drop_pct", "—"),
                    "section_tag":  c_sec,
                    "note":         note,
                    "is_reserve":   c.get("reserve", False),
                })
                seq += 1

    # ── Наружные сети ────────────────────────────────────────────────────────
    outdoor_results = project.get("_results", {}).get("outdoor_networks", [])
    for net in outdoor_results:
        cb = net.get("cable", {})
        if not cb.get("section_mm2"):
            continue
        records.append({
            "n":            seq,
            "cable_no":     "—",
            "mark":         cb.get("mark", "ВВГнг-LS"),
            "cores":        cb.get("cores", 4),
            "section":      cb.get("section_mm2"),
            "from_id":      net.get("panel_id", "ШУНО"),
            "from_name":    net.get("panel_name", "Щит наружного освещения"),
            "to_id":        net.get("id", "НО"),
            "to_name":      net.get("name", "Наружное освещение"),
            "length_plan":  _len_plan(cb),
            "length_calc":  _len_calc(cb),
            "install":      cb.get("install", "земля"),
            "breaker_a":    net.get("breaker", {}).get("rating", "—"),
            "voltage_drop": cb.get("voltage_drop_pct", "—"),
            "section_tag":  "НО",
            "note":         cb.get("routing_note", ""),
            "is_reserve":   False,
        })
        seq += 1

    return records


# ── DOCX ──────────────────────────────────────────────────────────────────────

_HEADERS = [
    "№", "№ КЛ", "Марка кабеля", "Жилы×Сечение",
    "Откуда", "Куда",
    "L план, м", "L расч., м",
    "Прокладка", "АВ, А", "ΔU, %", "Примечание",
]
_COL_W_CM = [0.7, 2.5, 3.2, 2.5, 3.5, 5.5, 2.0, 2.0, 3.0, 1.5, 1.5, 4.0]


def _set_cell(cell, text: str, size_pt: int = 9, bold: bool = False,
              center: bool = True):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    run.font.name  = "Times New Roman"


def _add_stamp(doc, proj: dict):
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc.add_paragraph()
    stamp = doc.add_table(rows=5, cols=4)
    stamp.style = "Table Grid"

    col_w   = [4.0, 5.5, 3.0, 2.5]
    hdr_row = ["Должность", "Фамилия И.О.", "Подпись", "Дата"]
    data    = [
        ("Разработал",    proj.get("designer",  "")),
        ("Проверил",      proj.get("checker",   "")),
        ("Нормоконтроль", proj.get("norm_head", "")),
        ("ГИП",           proj.get("gip",       "")),
    ]

    for cell, h, w in zip(stamp.rows[0].cells, hdr_row, col_w):
        from docx.shared import Cm
        cell.width = Cm(w)
        _set_cell(cell, h, bold=True)

    for i, (role, name) in enumerate(data, start=1):
        row = stamp.rows[i]
        _set_cell(row.cells[0], role, center=False)
        _set_cell(row.cells[1], name, center=False)


def generate_cable_journal(project: dict, docs_dir: Path) -> Path:
    """
    Генерирует кабельный журнал ГОСТ 21.613.
    Возвращает путь к созданному файлу (.docx или .txt).
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        _has_docx = True
    except ImportError:
        _has_docx = False

    docs_dir = Path(docs_dir)
    docs_dir.mkdir(parents=True, exist_ok=True)

    proj  = project.get("project", {})
    code  = proj.get("code", "ОБЪЕКТ")
    cables = _collect_cables(project)

    if not _has_docx:
        txt_path = docs_dir / f"{code}_cable_journal.txt"
        _write_txt(cables, proj, txt_path)
        return txt_path

    out_path = docs_dir / f"{code}_cable_journal.docx"
    doc = Document()

    # ── А3 альбом ─────────────────────────────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width    = Cm(42.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(1.0)
    sec.top_margin    = Cm(1.5)
    sec.bottom_margin = Cm(1.5)

    def _heading(text: str, size: int = 12, bold: bool = False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.name = "Times New Roman"

    _heading("КАБЕЛЬНЫЙ ЖУРНАЛ", size=13, bold=True)
    _heading(proj.get("name", ""), size=11)
    _heading(
        f"Код: {code}  |  Стадия: {proj.get('stage','Р')}  |  "
        f"Ред.: {proj.get('revision', 0)}  |  Дата: {proj.get('date','')}",
        size=10,
    )
    doc.add_paragraph()

    # ── Таблица ───────────────────────────────────────────────────────────────
    col_w = [Cm(w) for w in _COL_W_CM]

    table = doc.add_table(rows=1, cols=len(_HEADERS))
    table.style     = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_row = table.rows[0]
    for cell, h, w in zip(hdr_row.cells, _HEADERS, col_w):
        cell.width = w
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell(cell, h, bold=True, size_pt=9)

    # Строки данных
    current_section = None

    for rec in cables:
        # Заголовок раздела
        if rec["section_tag"] != current_section:
            current_section = rec["section_tag"]
            sep_row = table.add_row()
            merged = sep_row.cells[0].merge(sep_row.cells[len(_HEADERS) - 1])
            _set_cell(merged, f"  Раздел: {current_section}", bold=True,
                      size_pt=9, center=False)

        row = table.add_row()
        mark_section = f"{rec['cores']}×{rec['section']}"
        du = rec["voltage_drop"]
        du_str = f"{du:.2f}" if isinstance(du, float) else str(du)
        note = rec.get("note", "")
        if len(note) > 45:
            note = note[:42] + "..."

        values = [
            str(rec["n"]),
            rec["cable_no"],
            rec["mark"],
            mark_section,
            rec["from_id"],
            rec["to_name"][:40] if rec["to_name"] else rec["to_id"],
            str(int(rec["length_plan"])) if rec["length_plan"] else "—",
            str(int(rec["length_calc"])) if rec["length_calc"] else "—",
            rec["install"],
            str(rec["breaker_a"]),
            du_str,
            note,
        ]
        align_flags = [True, True, False, True, True, False,
                       True, True, False, True, True, False]

        for cell, val, w, center in zip(row.cells, values, col_w, align_flags):
            cell.width = w
            _set_cell(cell, val, size_pt=9, center=center)

    # ── Итог ─────────────────────────────────────────────────────────────────
    total_plan = sum(r["length_plan"] for r in cables)
    total_calc = sum(r["length_calc"] for r in cables)
    doc.add_paragraph()
    p_total = doc.add_paragraph()
    p_total.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rt = p_total.add_run(
        f"Итого кабельных линий: {len(cables)} шт.  |  "
        f"Длина по плану: {total_plan:.0f} м  |  "
        f"Расчётная длина (с запасом): {total_calc:.0f} м"
    )
    rt.font.size = Pt(10)
    rt.font.name = "Times New Roman"

    _add_stamp(doc, proj)

    doc.save(str(out_path))
    return out_path


# ── XLSX ──────────────────────────────────────────────────────────────────────

def generate_cable_journal_xlsx(project: dict, docs_dir: Path) -> Path:
    """
    Генерирует кабельный журнал в формате xlsx (ГОСТ 21.613).
    Рабочий документ для монтажа; штамп не включается.
    Возвращает путь к созданному файлу.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    docs_dir = Path(docs_dir)
    docs_dir.mkdir(parents=True, exist_ok=True)

    proj   = project.get("project", {})
    code   = proj.get("code", "ОБЪЕКТ")
    name   = proj.get("name", "")
    stage  = proj.get("stage", "Р")
    rev    = proj.get("revision", 0)
    date   = proj.get("date", "")
    cables = _collect_cables(project)

    out_path = docs_dir / f"{code}_cable_journal.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Кабельный журнал"

    # ── Стили ────────────────────────────────────────────────────────────────
    TNR = "Times New Roman"
    thin = Side(style="thin")
    border_all = Border(left=thin, right=thin, top=thin, bottom=thin)

    def _font(size=9, bold=False, italic=False, color="FF000000"):
        return Font(name=TNR, size=size, bold=bold, italic=italic, color=color)

    def _fill(hex_color):
        return PatternFill("solid", fgColor=hex_color)

    def _align(horizontal="left", wrap=True):
        return Alignment(horizontal=horizontal, vertical="center", wrap_text=wrap)

    fill_header  = _fill("FFCCCCCC")
    fill_section = _fill("FFD6E4F7")
    fill_du_warn = _fill("FFFFE066")

    # ── Ширины колонок ────────────────────────────────────────────────────────
    col_widths_ch = [4, 12, 14, 12, 16, 28, 10, 10, 14, 7, 7, 20]
    n_cols = len(_HEADERS)
    last_col = get_column_letter(n_cols)
    for i, w in enumerate(col_widths_ch, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w

    # ── Шапка документа (строки 1–4) ─────────────────────────────────────────
    def _header_row(row_num: int, text: str, size: int = 9, bold: bool = False):
        ws.merge_cells(f"A{row_num}:{last_col}{row_num}")
        cell = ws.cell(row=row_num, column=1, value=text)
        cell.font      = _font(size=size, bold=bold)
        cell.alignment = _align(horizontal="center")
        ws.row_dimensions[row_num].height = 18

    _header_row(1, "КАБЕЛЬНЫЙ ЖУРНАЛ  (ГОСТ 21.613)", size=13, bold=True)
    _header_row(2, name, size=11)
    _header_row(3,
        f"Код: {code}  |  Стадия: {stage}  |  Ред.: {rev}  |  Дата: {date}",
        size=9)
    ws.row_dimensions[4].height = 8  # пустой отступ

    # ── Строка заголовков (row 5) ─────────────────────────────────────────────
    for col_idx, h in enumerate(_HEADERS, start=1):
        cell = ws.cell(row=5, column=col_idx, value=h)
        cell.font      = _font(size=9, bold=True)
        cell.fill      = fill_header
        cell.alignment = _align(horizontal="center")
        cell.border    = border_all
    ws.row_dimensions[5].height = 25

    # freeze: 4 строки шапки + 1 строка заголовков → данные с 6-й строки
    ws.freeze_panes = "A6"

    # ── Заполнение данных ─────────────────────────────────────────────────────
    current_row    = 6
    current_section = None

    # Выравнивание по колонкам (центр / лево)
    align_flags = [True, True, False, True, True, False,
                   True, True, False, True, True, False]

    for rec in cables:
        # Строка-разделитель раздела
        if rec["section_tag"] != current_section:
            current_section = rec["section_tag"]
            ws.merge_cells(f"A{current_row}:{last_col}{current_row}")
            cell = ws.cell(row=current_row, column=1,
                           value=f"  Раздел: {current_section}")
            cell.font      = _font(size=9, bold=True)
            cell.fill      = fill_section
            cell.alignment = _align(horizontal="left")
            cell.border    = border_all
            ws.row_dimensions[current_row].height = 18
            current_row += 1

        is_reserve = rec.get("is_reserve", False)
        du = rec["voltage_drop"]
        du_str = f"{du:.2f}" if isinstance(du, float) else str(du)
        du_over = isinstance(du, float) and du > 5

        values = [
            str(rec["n"]),
            rec["cable_no"],
            rec["mark"],
            f"{rec['cores']}×{rec['section']}",
            rec["from_id"],
            rec["to_name"] or rec["to_id"],
            str(int(rec["length_plan"])) if rec["length_plan"] else "—",
            str(int(rec["length_calc"])) if rec["length_calc"] else "—",
            rec["install"],
            str(rec["breaker_a"]),
            du_str,
            rec.get("note", ""),
        ]

        for col_idx, (val, center) in enumerate(zip(values, align_flags), start=1):
            cell = ws.cell(row=current_row, column=col_idx, value=val)
            cell.alignment = _align(horizontal="center" if center else "left")
            cell.border    = border_all

            if is_reserve:
                cell.font = _font(size=9, italic=True, color="FF888888")
            else:
                cell.font = _font(size=9)

            # ΔU > 5% — жёлтый фон для колонки K (индекс 11)
            if col_idx == 11 and du_over:
                cell.fill = fill_du_warn
                cell.font = _font(size=9, bold=True,
                                  italic=is_reserve, color="FF000000")

        ws.row_dimensions[current_row].height = 15
        current_row += 1

    # ── Итоговая строка ───────────────────────────────────────────────────────
    total_plan = sum(r["length_plan"] for r in cables)
    total_calc = sum(r["length_calc"] for r in cables)
    current_row += 1  # пустая строка
    ws.merge_cells(f"A{current_row}:{last_col}{current_row}")
    summary_cell = ws.cell(row=current_row, column=1,
        value=(f"Итого кабельных линий: {len(cables)} шт.  |  "
               f"Длина по плану: {total_plan:.0f} м  |  "
               f"Расчётная длина (с запасом): {total_calc:.0f} м"))
    summary_cell.font      = _font(size=10, bold=True)
    summary_cell.alignment = _align(horizontal="left")
    ws.row_dimensions[current_row].height = 20

    wb.save(str(out_path))
    return out_path


# ── TXT fallback ──────────────────────────────────────────────────────────────

def _write_txt(cables: list[dict], proj: dict, path: Path):
    code = proj.get("code", "")
    lines = [
        "КАБЕЛЬНЫЙ ЖУРНАЛ (ГОСТ 21.613)",
        f"Объект: {proj.get('name','')} | Код: {code}",
        "",
        f"{'№':<4} {'№КЛ':<14} {'Марка':<14} {'Жил×Сеч':<8} "
        f"{'Откуда':<10} {'Куда':<30} {'Lпл,м':<7} {'Lрасч,м':<8} "
        f"{'Прокл.':<10} {'АВ,А':<6} {'ΔU,%':<6} Примечание",
        "-" * 130,
    ]
    current_section = None
    for rec in cables:
        if rec["section_tag"] != current_section:
            current_section = rec["section_tag"]
            lines.append(f"--- Раздел: {current_section} ---")
        du = rec["voltage_drop"]
        du_str = f"{du:.2f}" if isinstance(du, float) else str(du)
        to_name = (rec["to_name"] or rec["to_id"])[:28]
        lines.append(
            f"{rec['n']:<4} {rec['cable_no']:<14} {rec['mark']:<14} "
            f"{rec['cores']}×{rec['section']:<6} "
            f"{rec['from_id']:<10} {to_name:<30} "
            f"{rec['length_plan']!s:<7} {rec['length_calc']!s:<8} "
            f"{rec['install']:<10} {rec['breaker_a']!s:<6} {du_str:<6} "
            f"{rec.get('note','')}"
        )

    total_plan = sum(r["length_plan"] for r in cables)
    total_calc = sum(r["length_calc"] for r in cables)
    lines += [
        "",
        f"Итого: {len(cables)} линий  |  Lплан={total_plan:.0f}м  |  Lрасч={total_calc:.0f}м",
        "",
        f"Разработал: {proj.get('designer','')}",
        f"Проверил:   {proj.get('checker','')}",
    ]
    path.write_text("\n".join(lines), encoding="utf-8")
